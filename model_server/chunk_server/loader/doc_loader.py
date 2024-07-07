import docx
from langchain_core.documents import Document
from tqdm import tqdm


class DocLoader():
    def __init__(self, file_path):
        self.file_path = file_path

    def table_to_md(self, table):
        # Converts a docx table to markdown format
        md_table = []
        for row in table.rows:
            md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
            md_table.append(md_row)
        return '\n'.join(md_table)

    def load(self):
        docs = []
        try:
            document = docx.Document(self.file_path)
        except Exception as e:
            raise IOError(f"Error loading document: {e}")

        paragraphs = document.paragraphs

        # 段落解析：
        # 1. 以标题开始的段落为一个文档
        # 2. 以标题开始的段落之后的段落为该文档的内容
        # 3. 表格为一个文档
        current_doc = None
        for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
            if any(heading in paragraph.style.name for heading in ('Heading', '标题')):
                if current_doc is not None:
                    docs.append(Document(current_doc.strip()))
                current_doc = paragraph.text.strip() + "\n"  # Start a new document segment with the title
            else:
                if current_doc is not None:
                    current_doc += paragraph.text.strip() + "\n"
                else:
                    current_doc = paragraph.text.strip() + "\n"

        # Ensure the last current_doc is added
        if current_doc:
            docs.append(Document(current_doc.strip()))

        # Process tables
        tables = document.tables
        for table in tqdm(tables, desc=f"解析[{self.file_path}]的表格"):
            docs.append(Document(self.table_to_md(table), metadata={"format": "table"}))
        return docs
